from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    out_path = Path(__file__).resolve().parent / "ChatApp_Documentation.docx"

    doc = Document()

    doc.add_heading("VartalaP — Real‑Time Chat App (MERN + Socket.IO)", level=0)
    doc.add_paragraph(f"Project Documentation • Generated: {date.today().isoformat()}")

    add_heading(doc, "Project Summary", level=1)
    doc.add_paragraph(
        "VartalaP is a full‑stack real‑time chat application built with the MERN stack "
        "(MongoDB, Express, React, Node.js) and Socket.IO. It includes JWT authentication, "
        "online presence tracking, 1:1 messaging, message edit/delete, and file/image sharing "
        "with a responsive, modern UI."
    )

    add_heading(doc, "Core Features", level=1)
    add_bullets(
        doc,
        [
            "JWT Authentication: signup/login with hashed passwords (bcrypt) + protected routes",
            "Real‑time Presence: online users broadcast with Socket.IO",
            "1:1 Messaging: conversation auto-created on first message",
            "Edit Message (sender-only) with edited state",
            "Delete Message (sender-only) using soft-delete (keeps chat timeline)",
            "File/Image Sharing: up to 5 attachments, each up to 4MB; images preview inline",
            "Responsive UI: mobile-friendly layout that switches between Users and Chat",
            "Light/Dark theme toggle",
        ],
    )

    add_heading(doc, "Tech Stack", level=1)
    add_heading(doc, "Frontend", level=2)
    add_bullets(
        doc,
        [
            "React + Vite",
            "React Router",
            "Redux Toolkit",
            "Axios",
            "Tailwind CSS",
            "Socket.IO Client",
        ],
    )
    add_heading(doc, "Backend", level=2)
    add_bullets(
        doc,
        [
            "Node.js + Express",
            "Socket.IO",
            "MongoDB + Mongoose",
            "JWT + bcrypt",
            "dotenv + CORS",
        ],
    )

    add_heading(doc, "High‑Level Architecture", level=1)
    doc.add_paragraph(
        "Frontend (React) calls REST APIs for auth/history and uses Socket.IO for real‑time updates. "
        "Backend (Express) validates requests, persists data in MongoDB, and emits socket events "
        "for online presence and message updates."
    )

    add_heading(doc, "Backend APIs", level=1)
    doc.add_paragraph("Base URL: http://localhost:4000/api/v1")
    add_bullets(
        doc,
        [
            "POST /signup",
            "POST /login",
            "GET /getAllUsers (JWT required)",
            "POST /send-message (JWT required) — body: { receiverId, message, attachments? }",
            "GET /get-all-messages/:chatUserId (JWT required)",
            "PATCH /edit-message/:messageId (JWT required, sender-only) — body: { message }",
            "DELETE /delete-message/:messageId (JWT required, sender-only)",
        ],
    )

    add_heading(doc, "Socket.IO Events", level=1)
    add_bullets(
        doc,
        [
            "send-all-online-users (server → client): [userId, ...]",
            "new-message (server → client): message payload",
            "message-updated (server → client): updated message payload",
            "message-deleted (server → client): deleted message payload (soft-delete)",
        ],
    )

    add_heading(doc, "Database Models", level=1)
    add_bullets(
        doc,
        [
            "User: firstName, lastName, email, passwordHash, profilePicture",
            "Conversation: members[], messages[]",
            "Message: senderId, receiverId, message, attachments[], edited/deleted metadata",
        ],
    )

    add_heading(doc, "Local Setup", level=1)
    add_heading(doc, "Backend", level=2)
    add_bullets(
        doc,
        [
            "cd CheatChat/backend",
            "Copy .env.example → .env and set MONGODB_URI + JWT_SECRET",
            "npm install",
            "npm run dev",
        ],
    )
    add_heading(doc, "Frontend", level=2)
    add_bullets(
        doc,
        [
            "cd CheatChat/Frontend",
            "Copy .env.example → .env",
            "npm install",
            "npm run dev",
        ],
    )

    add_heading(doc, "Security & Git Hygiene", level=1)
    add_bullets(
        doc,
        [
            "Never commit .env files with real secrets; keep only .env.example in git.",
            "Rotate secrets immediately if you ever exposed credentials.",
            "Consider moving attachments to S3/Cloudinary/GridFS for production instead of storing base64 in MongoDB.",
        ],
    )

    add_heading(doc, "Resume Bullets (Copy/Paste)", level=1)
    add_bullets(
        doc,
        [
            "Built a real‑time MERN chat app with JWT auth, Socket.IO presence, and 1:1 messaging using React, Redux Toolkit, Express, and MongoDB.",
            "Designed Mongoose schemas and implemented protected REST APIs with validation and robust error handling.",
            "Implemented realtime delivery plus message edit/delete updates via Socket.IO.",
            "Added file/image sharing and a responsive mobile-first UI with Tailwind CSS.",
        ],
    )

    add_heading(doc, "Keywords / Hashtags", level=1)
    doc.add_paragraph(
        "Keywords: MERN Stack, Full Stack Developer, React, Redux Toolkit, Node.js, Express.js, MongoDB, "
        "Mongoose, Socket.IO, WebSockets, REST API, JWT, Authentication, Tailwind CSS, Vite, JavaScript."
    )
    doc.add_paragraph(
        "Hashtags: #FullStackDeveloper #MERNStack #ReactJS #ReduxToolkit #NodeJS #ExpressJS #MongoDB "
        "#SocketIO #WebSockets #RESTAPI #JWT #TailwindCSS #Vite #JavaScript #OpenToWork #Hiring"
    )

    doc.save(out_path)


if __name__ == "__main__":
    main()

